#!/usr/bin/python
# -*- coding: UTF-8 -*-

from docx import Document
from docx.shared import Inches

document = Document("emptySched.docx")

document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

document.add_heading('Heading, level 1', level=1)
# document.add_paragraph('Intense quote', style='IntenseQuote')

document.add_paragraph(
    'first item in unordered list', style='ListBullet'
)
document.add_paragraph(
    'first item in ordered list', style='ListNumber'
)

p = document.add_paragraph()
run = p.add_run("מה עם עברית?")
run.font.name = "Arial"
#run.font.rtl = True
run.add_text('\n')
run = p.add_run("עברית (ועוד עברית)?")
run.font.name = "Arial"
run.font.rtl = True
#document.add_paragraph()

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
#for item in recordset:
#    row_cells = table.add_row().cells
#    row_cells[0].text = str(item.qty)
#    row_cells[1].text = str(item.id)
#    row_cells[2].text = item.desc

#document.add_page_break()

document.save('demo.docx')