#!/usr/bin/python
# -*- coding: UTF-8 -*-
import itertools
import os
from docx import Document
from docx.enum.table import WD_TABLE_DIRECTION
from docx.shared import Pt

from abstract_base_helper import DocumentHelperInterface


class DocxHelper(DocumentHelperInterface):
    @staticmethod
    def set_cell(cell, value):
        cell.text = value['text']
        run = cell.paragraphs[0].runs[0]
        run.rtl = True
        font = run.font
        font.name = 'Arial'
        if 'size' in value:
            font.size = Pt(value['size'])
        else:
            font.size = Pt(14)

    def add_entry(self, cells, offset, name, time=None):
        if name is None:
            # Special case -- merge this cell with the preceding and following cell
            return
        if isinstance(name, str):
            name = {'text': name}
        if 'bold' in name and name['bold']:
            pass
        if time:
            if isinstance(time, str):
                time = {'text': time}
            self.set_cell(cells[0 + offset], name)
            self.set_cell(cells[1 + offset], time)
        else:
            self.set_cell(cells[0 + offset], name)
            cells[0 + offset].merge(cells[1 + offset])

    @staticmethod
    def create_table(worddoc, rows, cols):
        table = worddoc.add_table(rows=rows, cols=cols)
        table.table_direction = WD_TABLE_DIRECTION.RTL
        return table

    # For now, hardcoded to four columns
    def create_populate_table(self, worddoc, column1, column2):
        rows = max(len(column1), len(column2))
        table = self.create_table(worddoc, rows, 4)
        columns = itertools.zip_longest(column1, column2, fillvalue=None)
        row = 0
        for c1, c2 in columns:
            if c1 == None:
                c1 = (u'',)
            if c2 == None:
                c2 = (None,)
            if not (isinstance(c1, tuple) and isinstance(c2, tuple)):
                raise Exception("Error -- input is not tuple")
            cells = table.rows[row].cells
            row += 1
            self.add_entry(cells, 0, *c1)
            self.add_entry(cells, 2, *c2)

    @staticmethod
    def set_header(worddoc, data):
        paragraph = worddoc.add_paragraph()
        run = paragraph.add_run(data['text'])
        font = run.font
        # font.rtl = True  - TODO: This _should_ be a font setting, but it looks like that messes up other font parameters
        run.rtl = True
        font.name = 'Arial'
        if 'size' in data:
            font.size = Pt(data['size'])
        else:
            font.size = Pt(16)
        if 'bold' in data:
            font.bold = data['bold']
        else:
            font.bold = True
        if 'highlight' in data:
            font.highlight_color = data['highlight']

    @staticmethod
    def set_bullet_list(worddoc, data):
        for entry in data:
            paragraph = worddoc.add_paragraph(style='MyBullet')
            run = paragraph.add_run(entry)
            font = run.font
            font.rtl = True
            font.name = 'Arial'
            font.size = Pt(12)

    @staticmethod
    def create_doc(base_doc='emptySched.docx'):
        doc = Document(base_doc)
        return doc

    @staticmethod
    def save_doc(worddoc, month_name, year):
        worddoc.save(f'{os.getcwd()}{os.sep}{month_name}{year}.docx')

    @staticmethod
    def create_sheet():
        pass

    @staticmethod
    def set_row(sheet, row, date, name, shlishit=True):
        pass

    @staticmethod
    def save_sheet(sheet, year):
        pass
