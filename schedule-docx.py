	#!/usr/bin/python
# -*- coding: UTF-8 -*-

import times
import calendar
import sys
import datetime

from docx import Document
from docx.shared import Inches

def addEntry(cells, offset, name, time=None):
	if time:
		cells[0+offset].text = name
		cells[0+offset].paragraphs[0].runs[0].rtl = True
		cells[1+offset].text = time
		cells[1+offset].paragraphs[0].runs[0].rtl = True
	else:
		cells[0+offset].text = name # Need to figure out how to span multiple columns
		cells[0+offset].paragraphs[0].runs[0].rtl = True
		#row.td(name, colspan='2')

### 
# Need to add:
# Slichot
# Eruv Tavshilin
# DST switch
# Hoshana Rabah
# 9 Av special handling
# For Shabbat, first print Shabbat and date, then other holidays -- Chanuka, Rosh Chodesh, whatever

location = 'Israel'

year = int(sys.argv[1])
month = None
if len(sys.argv) > 2:
	month = int(sys.argv[2])

holidays = None
if month:
	holidays = calendar.getMonth(year, month, location)
else:
	holidays = calendar.getYear(year, location)

times.setLocation('Givat Zeev','Israel',31.86,35.17,'Asia/Jerusalem',0)

holidayDone = []

document = Document("emptySched.docx")

for jd in sorted(holidays):
	day = holidays[jd]
	dayTimes = times.getTimes(day['date'])
	
	dstActive = dayTimes['motzei'].dst().total_seconds() != 0
	
	gregDate = day['date'].strftime("%d.%m.%y")

	# Holiday types:
	# shabbat
	# chag
	# fast
	# RH (Rosh Hashanah)
	# YK (Yom Kippur)
	# CH (Chol Hamoed)
	# purim
	# SP (Shushan Purim)
	# erevPesach
	if 'RH' in day['type']:
		minchaErev = dayTimes['candleLighting'] + datetime.timedelta(minutes=5)
		shacharit = "07:30"
		dafYomi = "06:45"
		minchaK = dayTimes['motzei'] - datetime.timedelta(minutes=(70 + dayTimes['motzei'].minute % 5))
		
		# Do we say tashlich today?
		tashlich = (day['hebrew'][2] == 1 and dayTimes['candleLighting'].weekday() != 5) or (day['hebrew'][2] == 2 and dayTimes['candleLighting'].weekday() == 6)
		
		header = u"\n%s (%s - %s)" % (', '.join(a['hebrew'] for a in day['fullnames']), day['hebrewWritten'], gregDate)
		p = document.add_paragraph()
		r = p.add_run(header)
		r.rtl = True

		t = document.add_table(rows=1, cols=4)
		cells = t.rows[0].cells

		if day['hebrew'][2] == 2:
			addEntry(cells, 0, u"הדלקת נרות אחרי", dayTimes['motzei'].strftime("%H:%M"))
		else:
			addEntry(cells, 0, u"הדלקת נרות", dayTimes['candleLighting'].strftime("%H:%M"))
		
		if dayTimes['candleLighting'].weekday() != 5:
			addEntry(cells, 2, u"תקיעת שופר (משוער)", "09:30")

		cells = t.add_row().cells			
		if day['hebrew'][2] == 1:
			addEntry(cells, 0, u"מנחה וערבית", minchaErev.strftime("%H:%M"))
		else:
			addEntry(cells, 0, u"ערבית", dayTimes['motzei'].strftime("%H:%M"))
		if dayTimes['candleLighting'].weekday() != 5:
			addEntry(cells, 2, u"תקיעה שנייה לאחר התפילה")
		
		cells = t.add_row().cells			
		addEntry(cells, 0, u"שיעור בדף יומי", dafYomi)
		addEntry(cells, 2, u"חצות היום", dayTimes['noon'].strftime("%H:%M"))
		
		cells = t.add_row().cells			
		addEntry(cells, 0, u"שחרית", shacharit)
		if tashlich:
			addEntry(cells, 2, u"מנחה ותשליח", minchaK.strftime("%H:%M"))
		else:
			addEntry(cells, 2, u"מנחה", (minchaK + datetime.timedelta(minutes=20)).strftime("%H:%M"))
			
		# If it's not Erev Shabbat, add motzei chag line
		if day['hebrew'][2] == 2 and dayTimes['noon'].weekday() != 4:
			cells = t.add_row().cells			
			addEntry(cells, 2, u'ערבית ומוצ"ח', dayTimes['motzei'].strftime("%H:%M"))
		
	elif 'YK' in day['type']:
		minchaErev = dayTimes['candleLighting'] + datetime.timedelta(minutes=5)
		shacharit = "08:00"
		dafYomi = "07:15"
		minchaK = dayTimes['motzei'] - datetime.timedelta(minutes=(70 + dayTimes['motzei'].minute % 5))
		
		header = u"\n%s (%s - %s)" % (', '.join(a['hebrew'] for a in day['fullnames']), day['hebrewWritten'], gregDate)
		p = document.add_paragraph()
		r = p.add_run(header)
		r.rtl = True

		t = document.add_table(rows=1, cols=4)
		cells = t.rows[0].cells

		addEntry(cells, 0, u"הדלקת נרות", dayTimes['candleLighting'].strftime("%H:%M"))
		addEntry(cells, 2, u'יזכור (משוער)', "10:30")

		cells = t.add_row().cells			
		addEntry(cells, 0, u"כל נדרי וערבית", minchaErev.strftime("%H:%M"))
		addEntry(cells, 2, u"מנחה", (dayTimes['motzei'] - datetime.timedelta(hours=3, minutes=(dayTimes['motzei'].minute % 5))).strftime("%H:%M"))

		cells = t.add_row().cells			
		addEntry(cells, 0, u"שיעור אחרי שיר הייחוד")
		addEntry(cells, 2, u"נעילה (משוער)", (dayTimes['motzei'] - datetime.timedelta(minutes=(90 + dayTimes['motzei'].minute % 5))).strftime("%H:%M"))
		
		cells = t.add_row().cells			
		addEntry(cells, 0, u"שיעור בדף יומי", dafYomi)
		addEntry(cells, 2, u"ערבית ומוצאי חג", dayTimes['motzei'].strftime("%H:%M"))
		
		cells = t.add_row().cells			
		addEntry(cell, u"שחרית", shacharit)
		
	elif 'shabbat' in day['type'] or 'chag' in day['type']:
		# Mincha Erev Shabbat / Chag is 5 minutes after candle lighting
		minchaErev = dayTimes['candleLighting'] + datetime.timedelta(minutes=5)
		# Mincha Ktana is 1:10 before maariv, rounded back to 5 minutes
		minchaK = dayTimes['motzei'] - datetime.timedelta(minutes=(70 + dayTimes['motzei'].minute % 5))
		parentChildLearning = minchaK - datetime.timedelta(minutes=40)
		minchaG = "12:30"
		shacharit = "08:00"
		dafYomi = "07:15"
		if dstActive:
			minchaG = "13:30"
			shacharit = "08:30"
			dafYomi = "07:45"
		simchatTorah = any([name['english'] == 'Simchat Torah' for name in day['names']])
		yizkor = None
		if 'yizkor' in day:
			if location == 'Israel' and simchatTorah:
				yizkor = "11:45"
			else:
				yizkor = "10:00"
		
		lateMaariv = dayTimes['motzei'] + datetime.timedelta(minutes=35)
		dayBeforeIsChag = (jd - 1) in holidays and any([x in holidays[jd-1]['type'] for x in ('shabbat','chag','RH')])

		shabbatName = None
		otherName = []

		for fullname in day['fullnames']:
			if u'שבת' in fullname['hebrew']:
				shabbatName = fullname['hebrew']
			else:
				otherName.append(fullname['hebrew'])
		
		header = None
		if shabbatName:
			header = u"\n%s (%s - %s)" % (shabbatName, day['hebrewWritten'], gregDate)
			if len(otherName):
				header += u' - ' + u', '.join(a for a in otherName)
		else:
			header = u"\n%s (%s - %s)" % (', '.join(a['hebrew'] for a in day['fullnames']), day['hebrewWritten'], gregDate)

		p = document.add_paragraph()
		r = p.add_run(header)
		r.rtl = True

		t = document.add_table(rows=1, cols=4)
		cells = t.rows[0].cells

		if any([x in day['type'] for x in ('chag','RH')]) and dayBeforeIsChag:
			addEntry(cells, 0, u"הדלקת נרות אחרי", dayTimes['motzei'].strftime("%H:%M"))
		else:
			addEntry(cells, 0, u"הדלקת נרות", dayTimes['candleLighting'].strftime("%H:%M"))

		if yizkor:
			addEntry(cells, 2, u'יזכור (משוער)', yizkor)
		else:
			addEntry(cells, 2, u"מנחה גדולה", minchaG)
		
		name = ''
		if dayBeforeIsChag:
			if dayTimes['candleLighting'].weekday() == 5:
				name = u"קבלת שבת וערבית"
			else:
				name = u"ערבית"
		# TODO: Probably should calculate a different time for after mincha here...
		elif 'shabbat' in day['type']:
			name = u"מנחה וקבלת שבת"
		else:
			name = u"מנחה"
		cells = t.add_row().cells
		addEntry(cells, 0, name, minchaErev.strftime("%H:%M"))
		
		# If it's a day with Yizkor, we pushed off Mincha Gedolah, so do it now instead of the kids learning
		if yizkor:
			addEntry(cells, 2, u"מנחה גדולה", minchaG)
		elif 'shabbat' in day['type']:
			addEntry(cells, 2, u"לימוד הורים וילדים", parentChildLearning.strftime("%H:%M"))
		
		cells = t.add_row().cells
		addEntry(cells, 0, u"שיעור בדף יומי", dafYomi)
		if 'shabbat' in day['type']:
			name = u'מנחה קטנה וס"ש'
		else:
			name = u"מנחה קטנה"
		addEntry(cells, 2, name, minchaK.strftime("%H:%M"))
		
		cells = t.add_row().cells
		addEntry(cells, 0, u"שחרית", shacharit)
		if 'shabbat' in day['type']:
			addEntry(cells, 2, u'ערבית ומוצ"ש', '%s / %s' % (lateMaariv.strftime("%H:%M"), dayTimes['motzei'].strftime("%H:%M")))
		else:
			addEntry(cells, 2, u'ערבית ומוצ"ח', dayTimes['motzei'].strftime("%H:%M"))
		
		if 'shabbat' in day['type'] and day['mevarchim']:
			cells = t.add_row().cells
			addEntry(cells, 0, u"שיעור לנשים אחרי התפילה")
			cells[0].paragraphs[0].runs[0].bold = True
		
	elif 'RC' in day['type']:
		name = [x['hebrew'] for x in day['names'] if u'ראש' in x['hebrew']][0]
		if name not in holidayDone:
			holidayDone.append(name)
			desc = u'\n'
			twoDay = day['hebrew'][2] == 30
			if twoDay:
				desc += u'ימי ' + calendar.hebrewDayOfWeek(day['date'].weekday()) + u' - ' + calendar.hebrewDayOfWeek(day['date'].weekday()+1) + u' - '
				desc += name + u" "
				desc += holidays[jd+1]['date'].strftime("%d.%m.%y")
				desc += u' - ' + day['date'].strftime("%d.%m.%y")
			else:	
				desc += u'יום ' + calendar.hebrewDayOfWeek(day['date'].weekday()) + u' - '
				desc += name + u" "
				desc += day['date'].strftime("%d.%m.%y") 
			if day['date'].weekday() != 4:
				desc += u' - שחרית ב05:50'

			p = document.add_paragraph()
			r = p.add_run(desc)
			r.rtl = True

	elif 'fast' in day['type']:
		header = u"\n%s (%s - %s)" % (', '.join(a['hebrew'] for a in day['fullnames']), day['hebrewWritten'], gregDate)

		p = document.add_paragraph()
		r = p.add_run(header)
		r.rtl = True

		t = document.add_table(rows=1, cols=4)
		cells = t.rows[0].cells

		addEntry(cells, 0, u"תחילת הצום", dayTimes['fastBegins'].strftime("%H:%M"))
		addEntry(cells, 2, u"סוף הצום", dayTimes['fastEnds'].strftime("%H:%M"))
		
	elif 'purim' in day['type']:
		maariv = dayTimes['fastEnds'] - datetime.timedelta(minutes=(5 + dayTimes['motzei'].minute % 5))
		if dayTimes['candleLighting'].weekday() == 6:
			# If Purim is Motzei Shabbat, give 30 minutes before Maariv
			maariv = dayTimes['motzei'] + datetime.timedelta(minutes=(30 + dayTimes['motzei'].minute % 5))
		secondReading = maariv + datetime.timedelta(hours=2)
		
		header = u"\n%s (%s - %s)" % (', '.join(a['hebrew'] for a in day['fullnames']), day['hebrewWritten'], gregDate)

		p = document.add_paragraph()
		r = p.add_run(header)
		r.rtl = True

		t = document.add_table(rows=1, cols=4)
		cells = t.rows[0].cells

		addEntry(cells, 0, u"ערבית וקריאת מגילה", maariv.strftime("%H:%M"))
		addEntry(cells, 2, u"שחרית וקריאת מגילה", "07:00")
		
		cells = t.add_row().cells
		addEntry(cells, 0, u"קריאה שניה", secondReading.strftime("%H:%M"))
		addEntry(cells, 2, u"קריאה שניה", "09:30")

		cells = t.add_row().cells
		addEntry(cells, 0, u'')
		addEntry(cells, 2, u"מנחה", "13:00")

	elif 'chanuka' in day['type']:
		if 'chanuka' not in holidayDone:
			holidayDone.append('chanuka')
			chanuka1 = date_utils.calendar_util.hebrew_to_jd(year, 9, 25)
			chanuka8 = chanuka1 + 7
			
			day8 = holidays[chanuka8]

			desc = u'\n'
			desc += u"חנוכה, ימים " + calendar.hebrewDayOfWeek(day8['date'].weekday()) + '-' + calendar.hebrewDayOfWeek(day8['date'].weekday()) + ' '
			desc +=	u'אור לכ"ה כסלו ' + (day8['date']-datetime.timedelta(days=7)).strftime("%d.%m.%y") + u" עד " 
			desc +=	calendar.hebrewDayOfMonth(day8['hebrew'][2] - 1) + u"טבת "

			desc += day8['date'].strftime("%d.%m.%y")

			p = document.add_paragraph()
			r = p.add_run(desc)
			r.rtl = True

			t = document.add_table(rows=1, cols=4)
			cells = t.rows[0].cells

			addEntry(cells, 0, u"שחרית מנין א'", "05:50")
			addEntry(cells, 2, u"מנחה", (dayTimes['sunset'] - datetime.timedelta(minutes=(15 + dayTimes['sunset'].minute % 5))).strftime("%H:%M"))
			
			cells = t.add_row().cells
			addEntry(cells, 0, u"שחרית מנין ב'", "08:00")
			addEntry(cells, 2, u"ערבית", "20:10")
			
			for d in range(int(jd-0.5),int(chanuka8+0.5)):
				if holidays[d+0.5]['date'].weekday() == 4:
					# There's a Friday in Chanuka (well, duh, but specifically for the case of Tevet)
					cells = t.add_row().cells
					addEntry(cells, 0, u"שחרית מנין א' ביום ו'", "06:00")
		
	elif 'erevPesach' in day['type']:
		header = u"\n%s (%s - %s)" % (', '.join(a['hebrew'] for a in day['fullnames']), day['hebrewWritten'], gregDate)

		p = document.add_paragraph()
		r = p.add_run(header)
		r.rtl = True

		t = document.add_table(rows=1, cols=4)
		cells = t.rows[0].cells

		addEntry(cells, 0, u"סוף זמן אכילת חמץ", dayTimes['chametzEating'].strftime("%H:%M"))
		addEntry(cells, 2, u"סוף זמן שריפת חמץ", dayTimes['chametzBurning'].strftime("%H:%M"))

	elif 'rain' in day['type']:
		header = '\n'
		header += u"אור ל-ז' חשון בערבית "
		header += day['date'].strftime("(%d.%m.%y)")
		p = document.add_paragraph()
		r = p.add_run(header)
		r.rtl = True

		text = ' - '
		text += u'מתחילים לומר "ותן טל ומטר לברכה"'
		p = document.add_paragraph()
		r = p.add_run(text)
		r.rtl = True
		
	elif 'omer' in day['type'] and len(day['type']) == 1:
		pass
	else:
		header = "\n%s (%s - %s)" % (', '.join(a['hebrew'] for a in day['fullnames']), day['hebrewWritten'], gregDate)
		
		p = document.add_paragraph()
		r = p.add_run(header)
		r.rtl = True

	

document.save('schedule.docx')